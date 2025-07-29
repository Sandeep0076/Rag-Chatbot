import logging
import os
import re
import shutil
import subprocess
import tempfile
import time

from docx import Document


class TextExtractor:
    def __init__(self, configs=None):
        """Initialize TextExtractor with optional config."""
        self.configs = configs

    def extract_text_from_doc(self, file_path: str) -> str:
        """Main function to extract text from DOC/DOCX files."""
        logging.info(f"Extracting text from DOC/DOCX file: {file_path}")
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()

        try:
            if file_extension == ".docx":
                text = self._process_docx(file_path)
            elif file_extension == ".doc":
                text = self._process_doc(file_path)
            else:
                error_msg = f"Unsupported document extension: {file_extension}"
                logging.error(error_msg)
                return f"ERROR: {error_msg}"

            # Save extracted text for diagnosis
            self._save_extracted_text_for_diagnosis(file_path, text)

            # Check if extracted text is too short
            if len(text.strip()) < 100:
                return (
                    "ERROR: Unable to extract sufficient text from this file "
                    "(less than 100 characters). Please try using the 'Chat with Image' feature instead."
                )

            return text

        except Exception as e:
            error_msg = f"Failed to extract text from {file_extension} file: {str(e)}"
            logging.error(error_msg)
            return f"ERROR: {error_msg}"

    def _save_extracted_text_for_diagnosis(self, file_path: str, extracted_text: str):
        """Save extracted text to a diagnostic file for testing purposes."""
        # Check if diagnostic saving is enabled
        if not self.configs or not getattr(
            self.configs, "save_extracted_text_diagnostic", False
        ):
            return

        try:
            # Create diagnostic directory if it doesn't exist
            diagnostic_dir = "diagnostic_extracted_texts"
            os.makedirs(diagnostic_dir, exist_ok=True)

            # Generate filename based on original file
            original_filename = os.path.basename(file_path)
            base_name = os.path.splitext(original_filename)[0]
            timestamp = int(time.time())
            diagnostic_filename = f"{base_name}_extracted_text_{timestamp}.txt"
            diagnostic_path = os.path.join(diagnostic_dir, diagnostic_filename)

            # Save the extracted text
            with open(diagnostic_path, "w", encoding="utf-8") as f:
                f.write("=== EXTRACTED TEXT DIAGNOSTIC FILE ===\n")
                f.write(f"Original file: {file_path}\n")
                f.write(f"Extraction timestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Text length: {len(extracted_text)} characters\n")
                f.write(f"Word count: {len(extracted_text.split())} words\n")
                f.write(f"Line count: {len(extracted_text.splitlines())} lines\n")
                f.write(f"Contains markdown headers: {'#' in extracted_text}\n")
                f.write(
                    f"Contains markdown lists: {'* ' in extracted_text or '- ' in extracted_text}\n"
                )
                f.write(f"Contains markdown bold: {'**' in extracted_text}\n")
                f.write(f"Contains markdown italic: {'*' in extracted_text}\n")
                f.write(f"Contains markdown code blocks: {'```' in extracted_text}\n")
                f.write(
                    f"Contains markdown links: {'[' in extracted_text and '](' in extracted_text}\n"
                )
                f.write(f"Contains tables: {'|' in extracted_text}\n")
                newline_check = "\\n" in repr(extracted_text)
                tab_check = "\\t" in repr(extracted_text)
                f.write(f"Contains newlines: {newline_check}\n")
                f.write(f"Contains tabs: {tab_check}\n")
                f.write(f"First 200 characters: {repr(extracted_text[:200])}\n")
                f.write(f"Last 200 characters: {repr(extracted_text[-200:])}\n")
                f.write(f"\n{'=' * 50}\n")
                f.write("FULL EXTRACTED TEXT:\n")
                f.write(f"{'=' * 50}\n\n")
                f.write(extracted_text)
                f.write(f"\n\n{'=' * 50}\n")
                f.write("END OF EXTRACTED TEXT\n")
                f.write(f"{'=' * 50}\n")

            logging.info(f"Saved extracted text diagnostic file: {diagnostic_path}")

        except Exception as e:
            logging.warning(f"Failed to save extracted text diagnostic file: {str(e)}")

    def _process_docx(self, file_path: str) -> str:
        """Process .docx files using python-docx."""
        logging.info(f"Processing .docx file with python-docx: {file_path}")
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        if not text.strip():
            logging.info("No text found in paragraphs, trying to extract from tables")
            text = self._extract_text_from_tables(doc)

        return text

    def _extract_text_from_tables(self, doc) -> str:
        """Extract text from tables in a .docx file."""
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                table_texts.append(
                    " ".join([cell.text for cell in row.cells if cell.text])
                )
        return "\n".join(table_texts)

    def _process_doc(self, file_path: str) -> str:
        """Process .doc files using multiple approaches."""
        logging.info(f"Processing .doc file: {file_path}")
        temp_dir = None
        try:
            temp_dir = tempfile.mkdtemp()
            temp_docx = os.path.join(temp_dir, os.path.basename(file_path) + ".docx")

            # Attempt 1: Convert .doc to .docx using LibreOffice/OpenOffice
            if self._convert_to_docx(file_path, temp_dir, temp_docx):
                return self._process_docx(temp_docx)

            # Attempt 2: Try antiword
            text = self._use_antiword(file_path)
            if text:
                return text

            # Attempt 3: Try python-docx directly
            text = self._use_python_docx_directly(file_path)
            if text:
                return text

            # Attempt 4: Advanced binary extraction
            return self._advanced_binary_extraction(file_path)
        finally:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)

    def _convert_to_docx(self, file_path: str, temp_dir: str, temp_docx: str) -> bool:
        """Attempt to convert .doc to .docx using LibreOffice/OpenOffice."""
        libreoffice_paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
            "libreoffice",  # Linux
            "soffice",  # Windows/Linux alternative
            "openoffice",  # OpenOffice alternative
        ]
        for office_path in libreoffice_paths:
            try:
                logging.info(f"Attempting to convert .doc to .docx using {office_path}")
                result = subprocess.run(
                    [
                        office_path,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        temp_dir,
                        file_path,
                    ],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if result.returncode == 0 and os.path.exists(temp_docx):
                    logging.info(
                        f"Successfully converted .doc to .docx using {office_path}"
                    )
                    return True
            except (subprocess.SubprocessError, FileNotFoundError):
                continue
        return False

    def _use_antiword(self, file_path: str) -> str:
        """Attempt to extract text using antiword."""
        try:
            logging.info("Attempting to use antiword for .doc file")
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0:
                logging.info("Successfully extracted text using antiword")
                return result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            logging.warning("antiword not available or failed")
        return ""

    def _use_python_docx_directly(self, file_path: str) -> str:
        """Attempt to extract text using python-docx directly."""
        try:
            logging.info("Attempting to use python-docx directly for .doc file")
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            if text.strip():
                logging.info("Extracted text using python-docx directly")
                return text
        except Exception as docx_err:
            logging.warning(f"python-docx failed: {str(docx_err)}")
        return ""

    def _advanced_binary_extraction(self, file_path: str) -> str:
        """Attempt advanced binary extraction for .doc files."""
        try:
            logging.info("Attempting advanced binary extraction for .doc file")
            with open(file_path, "rb") as f:
                binary_content = f.read()

            raw_text = "".join(chr(b) if 32 <= b < 127 else " " for b in binary_content)
            potential_text_chunks = re.findall(
                r"[A-Z][^\s][^\s][^\s][\w\s\.,;:!?\-\'\"\(\)\[\]{}]{20,}[\.\?!]",
                raw_text,
            )
            word_chunks = re.findall(
                r"[\w]{3,}\s+[\w]{3,}\s+[\w]{3,}[\w\s\.,;:!?\-\'\"\(\)\[\]{}]{10,}",
                raw_text,
            )
            all_chunks = potential_text_chunks + word_chunks

            if all_chunks:
                cleaned_chunks = []
                for chunk in all_chunks:
                    chunk = re.sub(r"\s+", " ", chunk)
                    chunk = re.sub(r"[^\w\s\.,;:!?\-\'\"\(\)\[\]{}]", "", chunk)
                    chunk = chunk.strip()
                    if len(chunk) > 20:
                        cleaned_chunks.append(chunk)
                text = "\n".join(cleaned_chunks)
                logging.info("Extracted text using advanced binary extraction")
                return text
        except Exception as bin_err:
            logging.error(f"Advanced binary extraction failed: {str(bin_err)}")
        return (
            "ERROR: Failed to extract text from .doc file using all available methods"
        )
